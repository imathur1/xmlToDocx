import xml.etree.ElementTree as ET
from docx import Document
from docx.shared import Pt


class xmlToDocx(object):

    def __init__(self, xmlFile):
        self.tree = ET.parse(xmlFile)
        self.document = Document()

    def cleanText(self, string):
        newString = ""  # New string that will have plain text
        index = 0  # Keeps track of the indexes in the string
        tracker = 0  # Keeps track if the value in the string is between < and >
        source = False  # Returns true if the letters 'src' are found in the image tag
        picture = ""  # Will store the source of the picture
        math = False  # Returns true if there's a math tag
        latex = False  # Returns true if there's latex in the math tag

        for i in string:

            # If there's a math tag, look for the latex tag, otherwise ignore the value and skip to
            # the end of the loop
            if math == True and latex == False:
                if string[index] == '>' and string[index - 1] == '"' and string[index - 2] == 'x' and string[
                    index - 3] == 'e' and string[index - 4] == 't':
                    latex = True
                else:
                    index += 1
                    continue

                    # If it's the closing tag </annotation> make latex and math false, otherwise add the
            # latex to the newString and skip to the end of the loop
            elif latex == True:
                if string[index] == '<' and string[index + 1] == '/' and string[index + 2] == 'a':
                    latex = False
                    math = False
                else:
                    newString += i
                    index += 1
                    continue

            # Finds if the letters src=" are found in the string 
            if (string[index] == '"' and string[index - 1] == '=' and string[index - 2] == 'c' and string[
                index - 3] == 'r' and string[index - 4] == 's'):
                source = True

            # If the source is true, it adds the letters following the 'src=""' to the picture string
            elif source == True:

                # The end of the source is signified by '""', so this is when the image is actually added 
                # to the document                
                if i == '"':

                    # If the letters '%20' are in the string, it adds a space (' ') to the picture source
                    # For example 'Untitled%20picture.png' would change to 'Untitled picture.png'
                    picture = picture.replace('%20', ' ')

                    # The width and height are preset because finding the height and width for all of the
                    # different types of pictures was difficult if it wasn't specified

                    # The pictures also show up at the top of the question, not in the actual text
                    # However in the Word document the pictures can be resized and moved into the question
                    self.document.add_picture(picture, width=Pt(25), height=Pt(25))

                    picture = ""
                    source = False
                else:
                    picture += i

            # If the letter in the string is '<', it ignores all the text until it reaches '>'
            # This signifies the end of the tag attributes and the beginning of the actual text  
            if i == '<' or i == '{':

                # There are three exceptions, if the tag is <sup>, or superscript, a '^' is added
                # If the tag is <br> a space is added
                # If the tag is <math> it becomes true
                if (string[index + 1] == 's' and string[index + 2] == 'u' and string[index + 3] == 'p'):
                    newString += '^'
                elif (string[index + 1] == 'b' and string[index + 2] == 'r'):
                    newString += ' '
                elif (string[index + 1] == 'm' and string[index + 2] == 'a' and string[index + 3] == 't' and string[
                    index + 4] == 'h'):
                    math = True
                tracker = 1

            elif i == '>' or i == '}':

                # If the tag is <tr> it creates a new line, or row
                # If the tag is <td> it creates a new column with spaces
                if (string[index - 1] == 'r' and string[index - 2] == 't'):
                    newString += '\n'
                elif (string[index - 1] == 'd' and string[index - 2] == 't'):
                    newString += '       '
                tracker = 0
            else:
                if tracker == 0:
                    newString += i
            index += 1

        return newString

    def cleanCharacters(self, text):
        # Removes all the unicode characters and replaces them with their equivalant in general 
        # punctuation
        characters = {
            "&nbsp;": " ", '&#8211;': "-", '&#8217;': "'", '&#8203;': ". ",
            '&minus;': "-", '&quot;': '"', '&#945;': "α", '&#949;': "ε",
            '&#x03B2;': "β", '&#x005E;': "^", '&#x2211;': "∑", '&sum;': "∑",
            '&#x2212;': "-", '&#x2217;': "*", '\sqrt': '√', '\sum': '∑',
            '"version":"1.1","math":': "", r'\beta': 'β', '&alpha;': "α"
        }

        for i in characters:
            text = text.replace(i, characters[i])
        return text

    def convert(self):
        h = self.document.add_paragraph()
        h.add_run("Note: Pictures can be resized and moved around, and the tables can be reformatted. "
                  "In LaTeX, '_' means subscript, '\hat' means an accent, and  '"r'\frac'"' means a fraction.").bold = True

        previousTag = ""  # Remembers the previous element's tag
        previousTag2 = ""  # Remembers the element's tag before the previous element
        previousAttrib = ""  # Remembers the previous element's attributes
        previousAttrib2 = ""  # Remembers the element's attribute before the previous element

        counter = 1  # Keeps track of the answer order (A, B, C, D)
        counter2 = 1  # Keeps track of the answer choices
        feedbackCounter = 1  # Keeps track of the answer feedback

        answer = 0  # Finds the correct answer
        number = 1  # Numbers the questions
        newLines = 0  # Adds a questionString line after each set of questions and answers
        error = 0  # Fixes the issue with the first tag being 'None'

        for elem in self.tree.iter():

            # Resets after every question 
            if elem.tag == 'item':
                feedbackCounter = 1

            # Checks if the tag is 'setvar' and the value is > 0, meaning it's the right answer
            if elem.tag == 'setvar':
                if float(elem.text) > 0:
                    answer = counter2
                counter2 += 1

            # All questions and answers have a 'mattext' element tag
            if elem.tag == 'mattext':

                # If the element's 'texttype' attribute is 'text/plain', then it is a form of
                # feedback and not an answer
                if elem.attrib["texttype"] == "text/plain":
                    counter = 1  # Fixes the issue if there's only 2 answer choices

                    # Displays the correct answer, bolds it, and adds a questionString line between each
                    # set of questions and answers
                    if newLines == 0:
                        p = self.document.add_paragraph()
                        if answer == 1:
                            p.add_run("Correct Answer: A\n\n").bold = True
                        elif answer == 2:
                            p.add_run("Correct Answer: B\n\n").bold = True
                        elif answer == 3:
                            p.add_run("Correct Answer: C\n\n").bold = True
                        elif answer == 4:
                            p.add_run("Correct Answer: D\n\n").bold = True
                        counter2 = 1
                    newLines += 1
                else:
                    newLines = 0

                # All questions and answers have a 'meterial' element tag before the 'mattext' tag    
                if previousTag == 'material':

                    # Questions have a 'flow' tag around the 'material' tag around the 'mattext' tag
                    if previousTag2 == 'flow':

                        string = str(elem.text)  # Converts the text into a str type

                        questionString = self.cleanText(string)
                        questionString = self.cleanCharacters(questionString)

                        self.document.add_paragraph(str(number) + ". " + questionString)
                        number += 1

                    # Answers have the 'flow_mat' tag around the 'material' tag around the 'mattext' tag
                    elif previousTag2 == 'flow_mat':

                        string = str(elem.text)
                        answerString = self.cleanText(string)
                        answerString = self.cleanCharacters(answerString)

                        # The first 'mattext' tag that reaches this point has a value of 'None'
                        # beacuse it has no text, so the variable error ensures that nothing is
                        # written to the document until it equals 1, which is after the first iteration
                        if error == 1:
                            if counter == 1:
                                self.document.add_paragraph("A. " + answerString)
                            elif counter == 2:
                                self.document.add_paragraph("B. " + answerString)
                            elif counter == 3:
                                self.document.add_paragraph("C. " + answerString)
                            elif counter == 4:
                                self.document.add_paragraph("D. " + answerString)
                            counter += 1
                        error = 1

                        # Feedback has the 'itemfeedback' tag around the 'material' tag around the 'mattext' tag
                    elif previousTag2 == 'itemfeedback':

                        string = str(elem.text)
                        identification = previousAttrib2['ident']

                        # If there's no feedback then the string is 'None'
                        if 'None' != string:

                            feedbackString = self.cleanText(string)
                            feedbackString = self.cleanCharacters(feedbackString)

                            # If there's two underscores then the feedback is for the entire question,
                            # not just one specific answer
                            if identification.count('_') == 2:
                                self.document.add_paragraph("Feedback: " + feedbackString)
                                feedbackCounter -= 1
                            else:
                                if feedbackCounter == 1:
                                    self.document.add_paragraph("Feedback A: " + feedbackString)
                                elif feedbackCounter == 2:
                                    self.document.add_paragraph("Feedback B: " + feedbackString)
                                elif feedbackCounter == 3:
                                    self.document.add_paragraph("Feedback C: " + feedbackString)
                                elif feedbackCounter == 4:
                                    self.document.add_paragraph("Feedback D: " + feedbackString)

                        feedbackCounter += 1

            # Updates the 2nd previous element's tag to the previous element's tag and the 
            # previous element's tag to the current one 
            previousTag2 = previousTag
            previousTag = elem.tag

            # Updates the 2nd previous element's attribute to the previous element's attribute 
            # and the previous element's attributeto the current one 
            previousAttrib2 = previousAttrib
            previousAttrib = elem.attrib

        self.document.save('q&a.docx')


def main():
    converter = xmlToDocx('questiondb.xml')
    converter.convert()


if __name__ == '__main__':
    main()
